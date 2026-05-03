"""
report_service.py
─────────────────
Generates the final Microsoft Word (.docx) maintenance report using
python-docx.

Formatting goals
────────────────
• Professional, table-driven layout
• Dark header row for the main task table
• "Solved" in green, "Pending" in amber
• "Waiting for RM confirmation" highlighted in yellow
• All cells with borders
• Landscape orientation for readability
• Fit table to page width
"""

from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor, Inches


# ── Colour palette ─────────────────────────────────────────────────────────────
CLR_HEADER_BG   = RGBColor(0x1A, 0x2B, 0x4A)   # dark navy
CLR_HEADER_FG   = RGBColor(0xFF, 0xFF, 0xFF)   # white
CLR_SUBHDR_BG   = RGBColor(0xD9, 0xE1, 0xF2)   # light blue-grey
CLR_SUBHDR_FG   = RGBColor(0x1A, 0x2B, 0x4A)   # navy text
CLR_SOLVED      = RGBColor(0x17, 0x6B, 0x2D)   # forest green
CLR_PENDING     = RGBColor(0xC6, 0x5F, 0x00)   # amber
CLR_HIGHLIGHT   = RGBColor(0xFF, 0xFF, 0x99)    # pale yellow highlight bg
CLR_ROW_ALT     = RGBColor(0xF5, 0xF8, 0xFF)   # very light blue for alt rows

FONT_NAME = "Calibri"


# ── Public API ─────────────────────────────────────────────────────────────────

def generate_word_report(data: dict[str, Any], output_path: str) -> None:
    """
    Generate a formatted .docx report and save it to *output_path*.

    Parameters
    ----------
    data : dict
        supervisor, team, shift, time_range, date, tasks (list of dicts)
    output_path : str
        Full path where the .docx file will be written.
    """
    doc = Document()
    _set_landscape(doc)
    _set_margins(doc)
    _add_styles(doc)

    supervisor = data.get("supervisor", "Supervisor")
    team       = data.get("team", "Team")
    shift      = data.get("shift", "Overnight")
    time_range = data.get("time_range", "12 AM - 8 AM")
    date       = data.get("date", "")
    tasks      = data.get("tasks", [])

    # ── Opening salutation ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Dear {supervisor},")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(
        "Please find below the details of the RL tasks completed today:"
    )
    run2.font.name = FONT_NAME
    run2.font.size = Pt(11)

    # ── Report title ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("Red-Light Maintenance Tasks Report")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = CLR_HEADER_BG

    # ── Header info table ─────────────────────────────────────────────────────
    _add_header_table(doc, shift, time_range, date, team)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Main task table ───────────────────────────────────────────────────────
    _add_task_table(doc, tasks)

    # ── Closing ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    cr = closing.add_run("Best regards,")
    cr.font.name = FONT_NAME
    cr.font.size = Pt(11)

    team_p = doc.add_paragraph()
    tr = team_p.add_run(team)
    tr.bold = True
    tr.font.name = FONT_NAME
    tr.font.size = Pt(11)

    doc.save(output_path)


# ── Section / style helpers ───────────────────────────────────────────────────

def _set_landscape(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def _set_margins(doc: Document) -> None:
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(1.5))


def _add_styles(doc: Document) -> None:
    """Ensure the Normal style uses our base font."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)


# ── Header table (Shift / Time / Date / Team) ─────────────────────────────────

def _add_header_table(
    doc: Document,
    shift: str,
    time_range: str,
    date: str,
    team: str,
) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _set_table_width(table)

    headers = ["Shift", "Time", "Date", "Team"]
    values  = [shift, time_range, date, team]

    # Label row
    hdr_row = table.rows[0]
    for i, (h, v) in enumerate(zip(headers, values)):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]

        label_run = p.add_run(f"{h}:  ")
        label_run.bold = True
        label_run.font.name = FONT_NAME
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = CLR_HEADER_BG

        val_run = p.add_run(v)
        val_run.font.name = FONT_NAME
        val_run.font.size = Pt(10)

        _shade_cell(cell, "D9E1F2")


# ── Main task table ───────────────────────────────────────────────────────────

COLUMNS = [
    ("#",               Cm(0.8)),
    ("Task",            Cm(2.6)),
    ("Site ID",         Cm(2.6)),
    ("Approach",        Cm(1.8)),
    ("Problem",         Cm(4.0)),
    ("Vendor",          Cm(2.0)),
    ("SAP\nNotification", Cm(2.6)),
    ("Action Taken",    Cm(5.5)),
    ("Current\nStatus", Cm(2.2)),
    ("Comments",        Cm(4.5)),
]


def _add_task_table(doc: Document, tasks: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    _set_table_width(table)

    # ── Set column widths ─────────────────────────────────────────────────────
    for col_idx, (_, width) in enumerate(COLUMNS):
        for cell in table.columns[col_idx].cells:
            cell.width = width

    # ── Header row ────────────────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for col_idx, (col_name, _) in enumerate(COLUMNS):
        cell = hdr_row.cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col_name)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = CLR_HEADER_FG
        _shade_cell(cell, "1A2B4A")

    # ── Data rows ─────────────────────────────────────────────────────────────
    for row_idx, task in enumerate(tasks):
        row = table.add_row()
        # Alternate row shading
        bg = "F5F8FF" if row_idx % 2 == 1 else "FFFFFF"

        values = [
            str(task.get("row_num", row_idx + 1)),
            task.get("task", "Field Service"),
            task.get("site_id", "N/A"),
            task.get("approach", "N/A"),
            task.get("problem", ""),
            task.get("vendor", "N/A"),
            task.get("sap_notification", "N/A"),
            task.get("action_taken", ""),
            task.get("current_status", "Solved"),
            task.get("comments", "Waiting for RM confirmation"),
        ]

        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]

            # Status cell — coloured text
            if col_idx == 8:
                run = p.add_run(value)
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = Pt(9)
                run.font.color.rgb = (
                    CLR_SOLVED if value.lower() == "solved" else CLR_PENDING
                )
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Comments cell — yellow highlight for default comment
            elif col_idx == 9 and "waiting for rm" in value.lower():
                run = p.add_run(value)
                run.font.name = FONT_NAME
                run.font.size = Pt(9)
                _highlight_cell_yellow(cell)
            else:
                # Row number centred
                if col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(value)
                run.font.name = FONT_NAME
                run.font.size = Pt(9)
                _shade_cell(cell, bg)


# ── XML / cell formatting helpers ─────────────────────────────────────────────

def _set_table_width(table) -> None:
    """Make table span the full page width."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)


def _shade_cell(cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _highlight_cell_yellow(cell) -> None:
    """Set cell background to pale yellow."""
    _shade_cell(cell, "FFFF99")
